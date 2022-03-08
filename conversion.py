from ast import Sub
from operator import index
from random import randint, random, sample
from matplotlib import style
import pandas as pd
from docx import Document
from docx.shared import Inches
from itertools import combinations



letters = ['a', 'b', 'c', 'd', 'e']
ans_list = ('')

exercises = [[1, 2, 3], [4, 5, 6], [7, 3, 5], [1, 2, 7], [4, 6, 1]]

fileds = [
    ['Original', 'Original', 'Original'], 
    ['Original', 'Original', 'Original'],
    ['Original', 'Mirror', 'Mirror'],
    ['Mirror', 'Mirror', 'Mirror'],
    ['Mirror', 'Mirror', 'Original']
]

document = Document()
subsets = []
df = ''
    
def gen_problem(mod, sec, exercise):
    f_index = exercises.index(exercise)
    mods = df.loc[(df['Mod'] == mod)]
    secs = mods.loc[(df['Sec'] == sec)]
    problems = secs.loc[(df['Type'] == 'Opening')]
    problem = problems['Original'].iloc[0]
    document.add_paragraph(
        problem + ' Considere as seguintes afirmativas a esse respeito.', style='List Number'
    )

    field = fileds[f_index]

    sub_sentences = secs.loc[(df['Type'] == 'Sentence')]

    for index in range(1, 4):
        real_field = field[index - 1]
        option = sub_sentences.loc[(df['Exercise'] == exercise[index-1])][field[index - 1]].iloc[0]
        p1 = document.add_paragraph(f"{getRoman(index)}. {option}")
        p1.paragraph_format.left_indent = Inches(.25)
    p = document.add_paragraph('São corretas apenas as afirmativas:')

    keys = secs.loc[(df['Type'] == 'Key')]
    index, exact_ans = 0, []
    for i, key in keys[(df['Exercise'].isin(exercise))].iterrows():
        index = index + 1
        if (key[field[index - 1]] == 'Verdadeira.'):
            exact_ans.append(index)

    ans_index = subsets.index(exact_ans)

    answer(ans_index)
    document.add_paragraph('Gabarito: ')


    explanations = secs.loc[(df['Type'] == 'Explanation')]
    i = 0
    for index in exercise:
        print(index)
        key = keys.loc[(df['Exercise'] == index)][field[i]].iloc[0]
        explanation = explanations.loc[(df['Exercise'] == index)][field[i]].iloc[0]
        p2 = document.add_paragraph(f"{getRoman(index)}. é {key} porque {explanation}")
        p2.paragraph_format.left_indent = Inches(.25)
        i = i + 1

    document.add_paragraph(f'Referência: Módulo {mod}, Seção {sec}')

def getSubsets(): 
    ''' all possible answers. Only one called''' 
    stuff = [1, 2, 3]
    for L in range(0, len(stuff)+1):
        for subset in combinations(stuff, L):
            subsets.append(list(subset))
    
def answer(ans_index):
    ''' possible answers generate
    ex: a) I b) I e II ... 
    '''
    rnd = randint(0, 4)
    indexs = sample(range(0, 8), 5)
    for i in range(0, 5):
        # exact answer insert(if exact answer is contained in indexes, not run)
        if (rnd == i & ans_index not in indexs):
            solution = solution_str(subsets[ans_index])
            p3 = document.add_paragraph(f"{letters[i]}) {solution}")
            p3.paragraph_format.left_indent = Inches(.25)
        # regular answer insert
        else:
            solution = solution_str(subsets[indexs[i]])
            p3 = document.add_paragraph(f"{letters[i]}) {solution}")
            p3.paragraph_format.left_indent = Inches(.25)

def solution_str(tuple):
    '''list type answer to string type answer
        ex: [1, 2] => I e II
    '''
    n = len(tuple)
    result = ''

    if n == 0:
        result = 'Nenhuma delas'
    else:
        seperator = ''
        for i in range(0, n):
            if i == n-1:
                seperator = ' e '
            else: 
                seperator = ', '
            result = result + seperator + getRoman(tuple[i])
        result = result[2:]
    return result

def getRoman(number):
    # integer to Roman Number
	num = [1, 4, 5, 9, 10, 40, 50, 90, 100, 400, 500, 900, 1000]
	sym = ["I", "IV", "V", "IX", "X", "XL", "L", "XC", "C", "CD", "D", "CM", "M"]
	i = 12
	result = ''
	while number:
		div = number // num[i]
		number %= num[i]

		while div:
			result = result + sym[i]
			div -= 1
		i -= 1
	return result

if __name__ == "__main__":
    # Data Extraction
    df = pd.read_excel('Input2.xlsx')
    df.dropna(axis=1, how='all', inplace=True)
    df.dropna(axis=0, how='all', inplace=True)
    df.rename(columns={'Unnamed: 1': 'Mod', 'Unnamed: 2': 'Sec', 'Unnamed: 3': 'Exercise', 'Unnamed: 4': 'Type', 'Unnamed: 5': 'Original', 'Unnamed: 6': 'Mirror'}, inplace=True)
    df = df.iloc[1:]

    # Answer subsets
    getSubsets()

    # produce docx
    for mod in range(1, 5):
        for sec in range(1, 3):
            for exercise in exercises:
                gen_problem(mod, sec, exercise)

    document.save('demo.docx')